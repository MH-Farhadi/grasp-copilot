"""
Generate Word document with complete evaluation plan for PRIME paper.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path


def add_heading(doc, text, level):
    doc.add_heading(text, level=level)


def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
    doc.add_paragraph()  # spacing
    return table


def main():
    doc = Document()

    # Title
    title = doc.add_heading("PRIME: Evaluation Plan", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Complete evaluation plan for LLM-based executive for interactive manipulation planning.")
    doc.add_paragraph()

    # =========================================================================
    # SECTION A: Executive-only benchmark
    # =========================================================================
    add_heading(doc, "A. Executive-Only Benchmark (Primary LLM Evaluation)", 1)
    add_paragraph(doc, "Isolates LLM decision policy. No planner, no physics. Main evidence for novelty.", bold=True)

    add_heading(doc, "Dataset", 2)
    doc.add_paragraph("• Held-out llm_contract.jsonl (separate test run or held-out split)")
    doc.add_paragraph("• Report total examples evaluated (e.g., 2000–5000)")

    add_heading(doc, "Models to Compare (Main Table)", 2)
    add_table(doc,
        ["ID", "Model", "Kind", "Purpose"],
        [
            ["H1", "Heuristic executive (ask-if-ambiguous)", "rule", "Non-LLM anchor"],
            ["H2", "Heuristic executive (always-ask)", "rule", "Upper bound on burden"],
            ["Z1", "Qwen2.5-7B-Instruct zero-shot", "LLM", "Shows value of fine-tuning"],
            ["Z2", "Qwen2.5-3B-Instruct zero-shot", "LLM", "Scaling reference"],
            ["F1", "Qwen2.5-7B-Instruct fine-tuned", "LLM", "Main method (yours)"],
            ["F2", "Qwen2.5-3B-Instruct fine-tuned", "LLM", "Scaling ablation"],
            ["Z3", "Llama-3.1-8B-Instruct zero-shot", "LLM", "Cross-family (optional)"],
        ]
    )

    add_heading(doc, "Ablations (on F1, Second Table)", 2)
    add_table(doc,
        ["ID", "Ablation", "What it tests"],
        [
            ["A1", "Remove memory.past_dialogs", "Dialog history importance"],
            ["A2", "Remove memory.last_prompt.context", "Context-awareness importance"],
            ["A3", "Remove user_state.mode", "Mode-conditioning importance"],
            ["A4", "Remove memory.candidates", "Candidate constraint importance"],
            ["A5", "Train on unbalanced data", "Rebalancing importance"],
            ["A6", "Train on 10% / 30% data", "Sample efficiency"],
        ]
    )

    add_heading(doc, "Metrics", 2)

    add_paragraph(doc, "Primary Metrics (report in main table)", bold=True)
    add_table(doc,
        ["Metric", "Definition", "Why it matters"],
        [
            ["Tool accuracy", "pred.tool == gt.tool", "Core decision correctness"],
            ["Motion object accuracy", "pred.args.obj == gt.args.obj (when tool is motion)", "Target selection correctness"],
            ["Interact kind accuracy", "pred.args.kind == gt.args.kind (when tool is INTERACT)", "Dialog type correctness"],
            ["Schema-valid rate", "Output parses as valid JSON + matches tool schema", "Reliability (invalid = hard fail)"],
        ]
    )

    add_paragraph(doc, "Secondary Metrics (report in appendix or detailed table)", bold=True)
    add_table(doc,
        ["Metric", "Definition", "Why it matters"],
        [
            ["JSON-valid rate", "Output parses as JSON", "Raw generation reliability"],
            ["Strict exact match", "Full tool call match (ignoring INTERACT text)", "Upper-bound reference"],
            ["Tool confusion matrix", "Confusion among INTERACT / APPROACH / ALIGN_YAW", "Failure mode analysis"],
        ]
    )

    add_paragraph(doc, "Per-Context Breakdown (key for reviewer insight)", bold=True)
    add_table(doc,
        ["Context type", "Examples", "Report accuracy"],
        [
            ["no_context", "First turn, no prior prompt", "✓"],
            ["intent_gate_candidates", '"Are you trying to grasp one of these?" YES/NO', "✓"],
            ["candidate_choice", '"Which object?" menu', "✓"],
            ["confirm", '"Do you want me to approach X?" YES/NO', "✓"],
            ["mode_select", '"APPROACH or ALIGN_YAW?"', "✓"],
            ["anything_else", 'Recovery: "Is there anything else?"', "✓"],
            ["help", "Align-yaw help flow", "✓"],
        ]
    )

    add_heading(doc, "Outputs", 2)
    doc.add_paragraph("• summary_all.csv / summary_all.json (per-model metrics)")
    doc.add_paragraph("• mistakes_<model>.jsonl (error analysis)")
    doc.add_paragraph("• Per-context accuracy table (for paper)")

    # =========================================================================
    # SECTION B: Real-user study (EXPANDED)
    # =========================================================================
    doc.add_page_break()
    add_heading(doc, "B. Real-User Study (End-to-End Validation)", 1)
    add_paragraph(doc, "Proves the executive improves real assistive interaction with actual humans.", bold=True)

    # ----- Task Definition -----
    add_heading(doc, "B.1 Task Definition", 2)
    add_paragraph(doc, "Task: Single-Object Reach-to-Grasp", bold=True)
    doc.add_paragraph("The user must grasp a designated target object among distractors using the robot arm. "
                      "Assistance is ON-DEMAND: the user decides when (and whether) to request help from the assistant.")

    add_table(doc,
        ["Element", "Specification"],
        [
            ["Task type", "Single reach-to-grasp (one target per trial)"],
            ["Goal", "Grasp the indicated target object (e.g., 'Pick up the mug')"],
            ["Start state", "Gripper at home position, all objects visible on table"],
            ["Success condition", "Gripper closes on the correct target object"],
            ["Failure condition", "Timeout (e.g., 120s) OR max attempts exceeded OR user gives up"],
            ["Assistance trigger", "User presses 'Request Assistance' button (on-demand)"],
            ["Reset between trials", "Gripper returns to home; objects may be reshuffled"],
        ]
    )

    doc.add_paragraph("Rationale: Single-object grasping provides clean trial boundaries, unambiguous success/failure, "
                      "and direct attribution of metrics to one goal. Sequential or pick-and-place tasks add confounds.")

    # ----- Trial Structure -----
    add_heading(doc, "B.2 Trial Structure", 2)
    add_paragraph(doc, "One Trial = One Target Object", bold=True)

    add_table(doc,
        ["Aspect", "Value"],
        [
            ["Trials per condition per user", "4–6"],
            ["Conditions", "2 (Baseline executive vs LLM executive)"],
            ["Total trials per user", "8–12"],
            ["Target users", "10–12 participants"],
            ["Total trials per condition", "~50–70 (sufficient for paired tests)"],
        ]
    )

    add_paragraph(doc, "Trial Matrix Example (8 trials per user)", bold=True)
    add_table(doc,
        ["Trial #", "Condition", "Difficulty", "Scene"],
        [
            ["1", "A", "Easy", "Scene 1 (3 objects, target separated)"],
            ["2", "A", "Hard", "Scene 2 (6 objects, target among similar)"],
            ["3", "A", "Easy", "Scene 3"],
            ["4", "A", "Hard", "Scene 4"],
            ["5", "B", "Easy", "Scene 1 (same layout, different target)"],
            ["6", "B", "Hard", "Scene 2"],
            ["7", "B", "Easy", "Scene 3"],
            ["8", "B", "Hard", "Scene 4"],
        ]
    )

    doc.add_paragraph("• Counterbalance condition order: half of users do A→B, half do B→A")
    doc.add_paragraph("• Same scenes across conditions to control for scene difficulty; only target object changes")

    # ----- Difficulty Levels -----
    add_heading(doc, "B.3 Difficulty Levels", 2)
    add_table(doc,
        ["Difficulty", "# Objects", "Ambiguity", "Target Location"],
        [
            ["Easy", "2–3", "Target clearly separated from others", "Unoccluded, easy reach"],
            ["Hard", "5–8", "Target among similar objects (e.g., two cans, two boxes)", "Partially occluded or clustered"],
        ]
    )

    # ----- Experimental Design -----
    add_heading(doc, "B.4 Experimental Design", 2)
    add_table(doc,
        ["Aspect", "Specification"],
        [
            ["Design", "Within-subject (each user tries both conditions)"],
            ["Participants", "10–12 users (report demographics, any impairments if applicable)"],
            ["Conditions", "(A) Baseline executive (heuristic) vs (B) LLM executive"],
            ["Order", "Counterbalanced (half A→B, half B→A)"],
            ["Trials per condition", "4–6 (balanced easy/hard)"],
            ["Total trials per user", "8–12"],
            ["Practice trials", "1–2 familiarization trials before data collection (excluded from analysis)"],
        ]
    )

    # ----- Handling On-Demand Assistance Variability -----
    doc.add_page_break()
    add_heading(doc, "B.5 Handling On-Demand Assistance Variability", 2)
    add_paragraph(doc, "IMPORTANT: Because assistance is on-demand, users may request help at different times "
                       "(or not at all). This creates variability that must be addressed explicitly.", bold=True)

    add_paragraph(doc, "User Behavior Types", bold=True)
    add_table(doc,
        ["Type", "Behavior", "Risk"],
        [
            ["Never-askers", "Complete task solo, never request help", "Cannot evaluate LLM"],
            ["Early-askers", "Request help immediately or very early", "Inflates # interactions"],
            ["Late-askers", "Struggle first, ask only when stuck", "Time penalty not due to LLM"],
            ["Selective-askers", "Ask only in genuinely ambiguous situations", "Ideal but rare"],
        ]
    )

    add_paragraph(doc, "Solution 1: Log Assistance Request as a Dependent Variable", bold=True)
    doc.add_paragraph("Do NOT force users to request help. Instead, log WHEN and WHY they request:")
    add_table(doc,
        ["Logged Variable", "Definition"],
        [
            ["# assistance requests per trial", "How many times user pressed 'Help'"],
            ["Time to first assistance request (s)", "Seconds from trial start to first 'Help' press"],
            ["Task state at request", "Gripper pose vs target, # nearby candidates, ambiguity level"],
            ["Reason for request (optional)", "Post-trial self-report: 'Why did you ask for help?'"],
        ]
    )

    add_paragraph(doc, "Solution 2: Report Metrics Conditioned on Assistance Use", bold=True)
    doc.add_paragraph("Report metrics SEPARATELY for:")
    doc.add_paragraph("• All trials (overall system performance)")
    doc.add_paragraph("• Assisted trials only (where user requested help at least once)")
    doc.add_paragraph("• Unassisted trials only (user completed solo)")
    doc.add_paragraph("")
    doc.add_paragraph("This allows claims like: 'Among trials where users sought help, the LLM executive reduced "
                      "corrections by 40%.'")

    add_paragraph(doc, "Solution 3: Normalize by Assistance Opportunity", bold=True)
    add_table(doc,
        ["Metric", "Definition"],
        [
            ["Success rate per assistance request", "Did the request lead to successful completion?"],
            ["Replies per assistance session", "From first request to task completion or abandonment"],
            ["Corrections per assistance session", "NO / None-of-them within that session"],
            ["Time after first request", "Time from first 'Help' to completion (isolates LLM effect)"],
        ]
    )

    add_paragraph(doc, "Solution 4: Post-Hoc Stratification by Help-Seeking Style", bold=True)
    doc.add_paragraph("After data collection, classify users by their help-seeking tendency:")
    add_table(doc,
        ["Style", "Criteria"],
        [
            ["Low-reliance", "< 1 assistance request per trial on average"],
            ["High-reliance", "≥ 2 assistance requests per trial on average"],
        ]
    )
    doc.add_paragraph("Report metrics WITHIN each stratum to show LLM helps across user types.")

    add_paragraph(doc, "Solution 5: Optional Forced-Assistance Subset", bold=True)
    doc.add_paragraph("For 1–2 trials per user, AUTO-TRIGGER assistance at a predefined moment "
                      "(e.g., when gripper enters candidate zone). This removes user decision variability "
                      "and provides a clean comparison. Mark these trials separately in analysis.")

    add_paragraph(doc, "Solution 6: Subjective 'Right Time' Assessment", bold=True)
    doc.add_paragraph("In the post-trial questionnaire, ask:")
    doc.add_paragraph("• 'Did you feel you asked for help at the right time?' (1–7)")
    doc.add_paragraph("• 'Did the assistant intervene when you needed it?' (1–7)")
    doc.add_paragraph("• 'Did you feel the assistant was too pushy / too passive?' (1–7)")

    # ----- Metrics -----
    doc.add_page_break()
    add_heading(doc, "B.6 Metrics (Complete List)", 2)

    add_paragraph(doc, "A. Overall Trial Metrics (all trials)", bold=True)
    add_table(doc,
        ["Metric", "Definition", "Primary?"],
        [
            ["Task success rate", "1 if target grasped, 0 otherwise", "✓"],
            ["Time to completion (s)", "Start to successful grasp (or timeout)", "✓"],
            ["# assistance requests", "Times user pressed 'Help' during trial", "✓"],
            ["# task failures", "Did not complete within budget", "✓"],
        ]
    )

    add_paragraph(doc, "B. Assistance-Conditioned Metrics (assisted trials only)", bold=True)
    add_table(doc,
        ["Metric", "Definition", "Primary?"],
        [
            ["# replies per trial", "Total user inputs during assistance sessions", "✓ (key claim)"],
            ["# prompts (INTERACTs) per trial", "Assistant queries per trial", "✓"],
            ["# corrections per trial", "User said NO / None-of-them / rejected confirmation", "✓"],
            ["# wrong-target commits", "Motion toward non-intended object", "✓"],
            ["Time to first request (s)", "Start to first 'Help' press", "○"],
            ["Time after first request (s)", "First 'Help' to completion", "✓"],
            ["# recovery episodes", "Task succeeded after initial rejection", "○"],
        ]
    )

    add_paragraph(doc, "C. Per-Assistance-Session Metrics", bold=True)
    add_table(doc,
        ["Metric", "Definition", "Primary?"],
        [
            ["Replies per session", "Inputs from request to session end (success/abandon/timeout)", "✓"],
            ["Corrections per session", "NO / None-of-them within session", "✓"],
            ["Session success rate", "Did session lead to successful grasp?", "✓"],
            ["Session duration (s)", "Time from request to session end", "○"],
        ]
    )

    add_paragraph(doc, "D. Safety / Errors", bold=True)
    add_table(doc,
        ["Metric", "Definition", "Primary?"],
        [
            ["# collisions", "Robot-object or object-object contact (if tracked)", "○"],
            ["# safety stops", "E-stop or manual intervention", "○"],
            ["# invalid LLM outputs", "Assistant returned malformed JSON (if applicable)", "○"],
        ]
    )

    add_paragraph(doc, "E. Subjective (post-condition questionnaire)", bold=True)
    add_table(doc,
        ["Instrument", "Items", "Scale"],
        [
            ["NASA-TLX", "Mental, Physical, Temporal demand; Performance; Effort; Frustration", "0–100 (or 21-point)"],
            ["Custom Likert", "Helpfulness, Trust, Predictability, Annoyance (reverse)", "1–7"],
            ["Assistance Timing", "'Asked at right time?', 'Assistant too pushy/passive?'", "1–7"],
            ["Preference", "'Which assistant did you prefer?' (forced choice + open comment)", "Binary"],
        ]
    )

    # ----- What to Log -----
    add_heading(doc, "B.7 Data to Log Per Trial", 2)
    add_table(doc,
        ["Category", "Variables"],
        [
            ["Trial info", "trial_id, user_id, condition (A/B), difficulty (easy/hard), scene_id, target_object"],
            ["Timing", "trial_start_time, trial_end_time, time_to_completion"],
            ["Outcome", "success (0/1), failure_reason (timeout/user_abort/collision/other)"],
            ["Assistance requests", "List of {timestamp, gripper_pose, candidates_nearby, ambiguity_score}"],
            ["Assistance sessions", "List of {start_time, end_time, # replies, # corrections, outcome}"],
            ["User inputs", "Full log of all user inputs (button presses, menu selections)"],
            ["LLM outputs", "Full log of all assistant tool calls and raw outputs"],
            ["Errors", "# collisions, # safety_stops, # invalid_llm_outputs"],
        ]
    )

    # ----- Analysis -----
    add_heading(doc, "B.8 Analysis Plan", 2)

    add_paragraph(doc, "Primary Analysis: Paired Within-Subject", bold=True)
    doc.add_paragraph("Each user contributes data to both conditions → pair by user.")
    doc.add_paragraph("For each user:")
    doc.add_paragraph("  Δ_time = mean(time_LLM) − mean(time_baseline)")
    doc.add_paragraph("  Δ_replies = mean(replies_LLM) − mean(replies_baseline)")
    doc.add_paragraph("Test: Paired t-test on Δ values (or Wilcoxon signed-rank if non-normal)")

    add_paragraph(doc, "Difficulty Interaction: 2×2 Repeated-Measures ANOVA", bold=True)
    doc.add_paragraph("• Factor 1: Condition (Baseline vs LLM)")
    doc.add_paragraph("• Factor 2: Difficulty (Easy vs Hard)")
    doc.add_paragraph("• Test interaction: Does LLM help MORE in hard trials?")

    add_paragraph(doc, "Handling Unassisted Trials", bold=True)
    doc.add_paragraph("• For OVERALL metrics (success, time): include all trials")
    doc.add_paragraph("• For ASSISTANCE-SPECIFIC metrics (replies, corrections): report separately for assisted vs unassisted")
    doc.add_paragraph("• Sensitivity analysis: report both 'all trials' and 'assisted-only' versions")

    add_paragraph(doc, "Statistical Reporting", bold=True)
    add_table(doc,
        ["Comparison", "Test"],
        [
            ["Time, replies, TLX subscales", "Paired t-test or Wilcoxon signed-rank (if non-normal)"],
            ["Success rate", "McNemar's test"],
            ["Effect size", "Cohen's d (or r for non-parametric)"],
            ["Report format", "Mean ± 95% CI, p-value, effect size"],
        ]
    )

    # ----- Reporting Structure -----
    add_heading(doc, "B.9 Recommended Reporting Structure", 2)

    add_paragraph(doc, "Table 1: Overall Performance (all trials)", bold=True)
    add_table(doc,
        ["Metric", "Baseline (mean ± SD)", "LLM (mean ± SD)", "Δ", "p", "d"],
        [
            ["Success rate", "...", "...", "...", "...", "..."],
            ["Time to completion (s)", "...", "...", "...", "...", "..."],
            ["# assistance requests", "...", "...", "...", "...", "..."],
        ]
    )

    add_paragraph(doc, "Table 2: Assistance-Conditioned Metrics (assisted trials only)", bold=True)
    add_table(doc,
        ["Metric", "Baseline (mean ± SD)", "LLM (mean ± SD)", "Δ", "p", "d"],
        [
            ["Time after first request (s)", "...", "...", "...", "...", "..."],
            ["# replies per trial", "...", "...", "...", "...", "..."],
            ["# corrections per trial", "...", "...", "...", "...", "..."],
            ["Assisted success rate", "...", "...", "...", "...", "..."],
        ]
    )

    add_paragraph(doc, "Table 3: Stratified by User Help-Seeking Style (supplementary)", bold=True)
    add_table(doc,
        ["User style", "N", "LLM benefit (time Δ)", "LLM benefit (replies Δ)"],
        [
            ["Low-reliance", "...", "...", "..."],
            ["High-reliance", "...", "...", "..."],
        ]
    )

    add_paragraph(doc, "Table 4: NASA-TLX Subscales", bold=True)
    add_table(doc,
        ["Subscale", "Baseline (mean ± SD)", "LLM (mean ± SD)", "p", "d"],
        [
            ["Mental Demand", "...", "...", "...", "..."],
            ["Physical Demand", "...", "...", "...", "..."],
            ["Temporal Demand", "...", "...", "...", "..."],
            ["Performance", "...", "...", "...", "..."],
            ["Effort", "...", "...", "...", "..."],
            ["Frustration", "...", "...", "...", "..."],
        ]
    )

    # ----- Key Figures -----
    add_heading(doc, "B.10 Key Figures for Paper", 2)
    doc.add_paragraph("• Bar chart: # replies per trial (baseline vs LLM), split by easy/hard")
    doc.add_paragraph("• Bar chart: Time to completion (baseline vs LLM)")
    doc.add_paragraph("• Grouped bar or radar: NASA-TLX subscales comparison")
    doc.add_paragraph("• Scatter/box: Time to first assistance request vs task difficulty")
    doc.add_paragraph("• Pie/stacked bar: Failure attribution breakdown (if failures occur)")

    # =========================================================================
    # SECTION C: Closed-loop simulation
    # =========================================================================
    doc.add_page_break()
    add_heading(doc, "C. Closed-Loop Simulation (Supporting Evidence)", 1)
    add_paragraph(doc, "Controlled stress tests + failure attribution. Main paper: 1 figure/table. Appendix: full sweeps.", bold=True)

    add_heading(doc, "Conditions", 2)
    add_table(doc,
        ["Condition", "Objects", "Ambiguity", "Noise"],
        [
            ["Easy", "2–3", "Low (separated)", "None"],
            ["Hard", "5–10", "High (clustered, similar labels)", "Optional: pose noise, occlusion"],
        ]
    )

    add_heading(doc, "Methods", 2)
    add_table(doc,
        ["ID", "Method", "Notes"],
        [
            ["LLM", "Qwen2.5-7B fine-tuned", "Main method"],
            ["H1", "Heuristic (ask-if-ambiguous)", "Baseline"],
            ["H2", "Always-ask", "Upper bound on burden"],
        ]
    )

    add_heading(doc, "User Model (Simulated)", 2)
    doc.add_paragraph("• Scripted replies: YES/NO, object label selection, 'None of them'")
    doc.add_paragraph("• Optional noise: hesitation, random mode switches, reply errors")

    add_heading(doc, "Episodes", 2)
    doc.add_paragraph("• 200–500 episodes per condition × method (report seed)")

    add_heading(doc, "Metrics", 2)

    add_paragraph(doc, "Executive-Centric (planner-independent)", bold=True)
    add_table(doc,
        ["Metric", "Definition", "Why it matters"],
        [
            ["Correct commit rate", "Fraction of motion commits targeting intended object", "Executive accuracy in context"],
            ["Query efficiency", "# INTERACTs per correct commit", "Minimal-input claim"],
            ["Decision success rate", "Reached 'ready-to-grasp' state for intended object (ignoring planner feasibility)", "Executive goal achievement"],
        ]
    )

    add_paragraph(doc, "End-to-End (with failure attribution)", bold=True)
    add_table(doc,
        ["Metric", "Definition", "Why it matters"],
        [
            ["E2E success rate", "Grasp completed", "Overall system performance"],
            ["Failure breakdown", "Categorize each failure: {executive wrong decision, invalid output, planner failure, perception failure}", "Attribution"],
        ]
    )

    add_paragraph(doc, "User Burden", bold=True)
    add_table(doc,
        ["Metric", "Definition", "Why it matters"],
        [
            ["# replies per episode", "Total user inputs", "Key claim"],
            ["# replies per success", "Only for successful episodes", "Efficiency"],
            ["# corrections per episode", "NO / None-of-them events", "Annoyance proxy"],
        ]
    )

    add_paragraph(doc, "Efficiency", bold=True)
    add_table(doc,
        ["Metric", "Definition", "Why it matters"],
        [
            ["Steps to success", "Total discrete steps (motions + interactions)", "Episode length"],
            ["Time to success", "Simulated or wall-clock time", "Real-time relevance"],
        ]
    )

    add_heading(doc, "Key Figure (Main Paper)", 2)
    doc.add_paragraph("Pareto plot: X = avg user replies per success, Y = success rate")
    doc.add_paragraph("• Each point = one method; show easy vs hard as separate marker shapes")
    doc.add_paragraph("• Shows 'minimal input' tradeoff clearly")

    add_heading(doc, "Additional Figures (Appendix)", 2)
    doc.add_paragraph("• Failure attribution stacked bar (per method × condition)")
    doc.add_paragraph("• Per-context accuracy heatmap (method × context type)")
    doc.add_paragraph("• Scaling curve (if you vary clutter levels beyond easy/hard)")

    # =========================================================================
    # Summary
    # =========================================================================
    doc.add_page_break()
    add_heading(doc, "D. Summary: Paper Structure", 1)

    add_table(doc,
        ["Section", "Content", "Key Metrics"],
        [
            ["Experiments A: Executive benchmark", "Models + ablations on contract JSONL", "Tool acc, motion obj acc, interact kind acc, per-context"],
            ["Experiments B: User study", "Real users, baseline vs LLM, on-demand assistance", "Success, time, replies (conditioned), corrections, NASA-TLX"],
            ["Experiments C: Simulation (1 fig/table)", "Easy vs hard, Pareto plot", "Success vs replies, failure attribution"],
            ["Appendix", "Full ablations, per-context tables, stratified user analysis, simulation sweeps", "All secondary metrics"],
        ]
    )

    # =========================================================================
    # Checklist
    # =========================================================================
    add_heading(doc, "E. Checklist Before Running Experiments", 1)
    checklist = [
        "Finalize held-out test contract JSONL (or separate run)",
        "Train Qwen2.5-3B fine-tuned checkpoint (for scaling ablation)",
        "Prepare ablation checkpoints (remove memory fields, retrain)",
        "IRB approval for user study (if required)",
        "Define task scenes + object layouts for easy/hard (4 scenes minimum)",
        "Implement 'Request Assistance' button and logging system",
        "Prepare questionnaires (NASA-TLX, Likert, Assistance Timing)",
        "Set up logging for ALL metrics (timestamps, poses, inputs, LLM outputs)",
        "Pilot test with 1–2 users to validate procedure",
        "Prepare counterbalancing schedule (user → condition order)",
    ]
    for item in checklist:
        doc.add_paragraph(f"☐ {item}")

    # Save
    out_path = Path(__file__).parent / "PRIME_Evaluation_Plan.docx"
    doc.save(out_path)
    print(f"[generate_eval_plan_doc] Saved: {out_path}")
    return out_path


if __name__ == "__main__":
    main()
